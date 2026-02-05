"""
Document Parser Service
Extracts text content from various document formats (PDF, DOCX, HTML, TXT)
"""
from pathlib import Path
from typing import Dict, Any, Optional
import structlog

# PDF parsing
try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

# DOCX parsing
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# HTML parsing
try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False


logger = structlog.get_logger(__name__)


class DocumentParserError(Exception):
    """Custom exception for document parsing errors"""
    pass


class DocumentParser:
    """
    Document parser supporting multiple file formats
    
    Supported formats:
        - PDF (.pdf)
        - Word (.docx)
        - HTML (.html)
        - Plain text (.txt)
    """
    
    def __init__(self):
        self.parsers = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.html': self._parse_html,
            '.htm': self._parse_html,
            '.txt': self._parse_txt,
        }
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse document and extract text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing:
                - text: Extracted text content
                - metadata: Additional parsing metadata
                
        Raises:
            DocumentParserError: If parsing fails
        """
        path = Path(file_path)
        
        if not path.exists():
            raise DocumentParserError(f"File not found: {file_path}")
        
        file_ext = path.suffix.lower()
        
        if file_ext not in self.parsers:
            raise DocumentParserError(f"Unsupported file type: {file_ext}")
        
        logger.info("parsing_document", file_path=str(path), file_type=file_ext)
        
        try:
            parser_func = self.parsers[file_ext]
            result = parser_func(path)
            
            # Add common metadata
            result['metadata']['file_size'] = path.stat().st_size
            result['metadata']['file_type'] = file_ext
            
            # Calculate text statistics
            text_length = len(result['text'])
            word_count = len(result['text'].split())
            
            result['metadata']['text_length'] = text_length
            result['metadata']['word_count'] = word_count
            
            logger.info("parsing_completed", 
                       file_path=str(path), 
                       text_length=text_length,
                       word_count=word_count)
            
            return result
        
        except Exception as e:
            logger.error("parsing_failed", file_path=str(path), error=str(e), exc_info=True)
            raise DocumentParserError(f"Failed to parse document: {str(e)}") from e
    
    def _parse_pdf(self, path: Path) -> Dict[str, Any]:
        """
        Parse PDF document using pypdf
        
        Args:
            path: Path to PDF file
            
        Returns:
            Dict with text and metadata
        """
        if not HAS_PYPDF:
            raise DocumentParserError("pypdf library not installed. Cannot parse PDF files.")
        
        try:
            text_parts = []
            metadata = {}
            
            with open(path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata['pdf_metadata'] = {
                        'author': pdf_reader.metadata.get('/Author'),
                        'creator': pdf_reader.metadata.get('/Creator'),
                        'producer': pdf_reader.metadata.get('/Producer'),
                        'subject': pdf_reader.metadata.get('/Subject'),
                        'title': pdf_reader.metadata.get('/Title'),
                    }
                
                metadata['page_count'] = len(pdf_reader.pages)
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            # Add page marker for citation tracking
                            text_parts.append(f"\n[PAGE {page_num}]\n{page_text}")
                    except Exception as e:
                        logger.warning("page_extraction_failed", 
                                     page_num=page_num, 
                                     error=str(e))
                        continue
            
            text = "\n".join(text_parts).strip()
            
            if not text:
                raise DocumentParserError("No text content extracted from PDF")
            
            return {'text': text, 'metadata': metadata}
        
        except pypdf.errors.PdfReadError as e:
            raise DocumentParserError(f"Invalid or corrupted PDF file: {str(e)}") from e
    
    def _parse_docx(self, path: Path) -> Dict[str, Any]:
        """
        Parse DOCX document using python-docx
        
        Args:
            path: Path to DOCX file
            
        Returns:
            Dict with text and metadata
        """
        if not HAS_DOCX:
            raise DocumentParserError("python-docx library not installed. Cannot parse DOCX files.")
        
        try:
            doc = DocxDocument(path)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Combine all text
            all_text = paragraphs + table_texts
            text = "\n".join(all_text).strip()
            
            if not text:
                raise DocumentParserError("No text content extracted from DOCX")
            
            # Metadata
            metadata = {
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
            }
            
            # Core properties if available
            if doc.core_properties:
                metadata['docx_metadata'] = {
                    'author': doc.core_properties.author,
                    'created': str(doc.core_properties.created) if doc.core_properties.created else None,
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else None,
                    'title': doc.core_properties.title,
                    'subject': doc.core_properties.subject,
                }
            
            return {'text': text, 'metadata': metadata}
        
        except Exception as e:
            raise DocumentParserError(f"Failed to parse DOCX: {str(e)}") from e
    
    def _parse_html(self, path: Path) -> Dict[str, Any]:
        """
        Parse HTML document using BeautifulSoup
        
        Args:
            path: Path to HTML file
            
        Returns:
            Dict with text and metadata
        """
        if not HAS_BS4:
            raise DocumentParserError("beautifulsoup4 library not installed. Cannot parse HTML files.")
        
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(['script', 'style', 'meta', 'link']):
                script.decompose()
            
            # Extract text
            text = soup.get_text(separator='\n', strip=True)
            
            if not text:
                raise DocumentParserError("No text content extracted from HTML")
            
            # Metadata
            metadata = {}
            
            # Try to get title
            title_tag = soup.find('title')
            if title_tag:
                metadata['title'] = title_tag.get_text(strip=True)
            
            # Count structural elements
            metadata['html_structure'] = {
                'headings': len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])),
                'paragraphs': len(soup.find_all('p')),
                'tables': len(soup.find_all('table')),
                'lists': len(soup.find_all(['ul', 'ol'])),
            }
            
            return {'text': text, 'metadata': metadata}
        
        except Exception as e:
            raise DocumentParserError(f"Failed to parse HTML: {str(e)}") from e
    
    def _parse_txt(self, path: Path) -> Dict[str, Any]:
        """
        Parse plain text document
        
        Args:
            path: Path to text file
            
        Returns:
            Dict with text and metadata
        """
        try:
            # Try UTF-8 first, fallback to latin-1
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            encoding_used = None
            
            for encoding in encodings:
                try:
                    with open(path, 'r', encoding=encoding) as file:
                        text = file.read()
                    encoding_used = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise DocumentParserError("Failed to decode text file with any supported encoding")
            
            text = text.strip()
            
            if not text:
                raise DocumentParserError("No text content in file")
            
            metadata = {
                'encoding': encoding_used,
                'line_count': len(text.split('\n')),
            }
            
            return {'text': text, 'metadata': metadata}
        
        except Exception as e:
            raise DocumentParserError(f"Failed to parse text file: {str(e)}") from e
